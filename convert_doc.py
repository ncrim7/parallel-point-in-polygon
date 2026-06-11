import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def parse_markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return False
        
    doc = Document()
    
    # Page Setup: 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    in_table = False
    table_lines = []
    
    i = 0
    num_lines = len(lines)
    
    while i < num_lines:
        line = lines[i]
        stripped = line.strip()
        
        # Handle Code Blocks
        if stripped.startswith("```"):
            if in_code_block:
                # Close code block
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                # Gray border / shading look
                run = p.add_run("\n".join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\n'))
            i += 1
            continue
            
        # Handle Tables
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # Table has ended
            in_table = False
            # Process table_lines
            parse_and_add_table(doc, table_lines)
            table_lines = []
            # Do not increment i, let current line be parsed normally
            
        # Skip empty lines (handled by spacing)
        if not stripped:
            i += 1
            continue
            
        # Handle Headings
        if stripped.startswith("#"):
            match = re.match(r'^(#+)\s+(.*)$', stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                
                run = p.add_run(text)
                run.bold = True
                
                if level == 1:
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50) # Dark Blue
                    p.paragraph_format.space_before = Pt(18)
                elif level == 2:
                    run.font.size = Pt(15)
                    run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E) # Soft Blue/Gray
                    p.paragraph_format.space_before = Pt(14)
                else:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            i += 1
            continue
            
        # Handle Images
        # Format: ![alt](path) or <img src="path" ... />
        img_match = re.search(r'!\[.*?\]\((.*?)\)', stripped)
        if not img_match:
            # HTML style
            img_match = re.search(r'<img\s+src=["\'](.*?)["\']', stripped)
            
        if img_match:
            img_path = img_match.group(1).strip()
            if os.path.exists(img_path):
                # Center the image
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                
                # Add picture (limit width to 5 inches to fit margins)
                p.add_run().add_picture(img_path, width=Inches(4.5))
                
                # Add a caption paragraph below
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(10)
                run_cap = p_cap.add_run("Grafik: Ölçüm ve Görselleştirme Sonuçları")
                run_cap.font.italic = True
                run_cap.font.size = Pt(9.5)
            else:
                print(f"Warning: Image path not found: {img_path}")
            i += 1
            continue
            
        # Handle Bullet Lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, bullet_text)
            i += 1
            continue
            
        # Handle Normal Paragraphs
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)
        i += 1
        
    # If file ends while in table
    if table_lines:
        parse_and_add_table(doc, table_lines)
        
    try:
        doc.save(docx_path)
        print(f"Successfully converted report to {docx_path}")
    except PermissionError:
        fallback_path = docx_path.replace(".docx", "_yeni.docx")
        doc.save(fallback_path)
        print(f"Permission Denied for {docx_path} (likely open in Word). Saved to {fallback_path} instead.")
    return True

def parse_inline_formatting(paragraph, text):
    # This function handles simple markdown inline formatting:
    # **bold** -> bold run
    # `code` -> monospace font run
    # $$math$$ or $math$ -> italic run
    
    # We will tokenise using regex
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`|\$.*?\$)', text)
    for token in tokens:
        if not token:
            continue
            
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xD4, 0x3F, 0x3A) # slightly red code
        elif token.startswith("$") and token.endswith("$"):
            # Inline math
            math_text = token.replace("$", "")
            run = paragraph.add_run(math_text)
            run.font.italic = True
        else:
            paragraph.add_run(token)

def parse_and_add_table(doc, table_lines):
    if len(table_lines) < 2:
        return
        
    # Parse header
    headers = [col.strip() for col in table_lines[0].split("|")[1:-1]]
    
    # Skip separator line (index 1)
    rows_data = []
    for line in table_lines[2:]:
        cols = [col.strip() for col in line.split("|")[1:-1]]
        rows_data.append(cols)
        
    num_rows = len(rows_data) + 1
    num_cols = len(headers)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        hdr_cells[col_idx].text = text
        set_cell_background(hdr_cells[col_idx], "2C3E50") # Dark Blue
        
        # Style Header Text
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White
        run.font.size = Pt(10)
        
    # Format Body Rows
    for row_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[row_idx + 1].cells
        
        # Alternating background shading
        bg_color = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"
        
        for col_idx, val in enumerate(row_data):
            # Safe boundary check
            if col_idx < len(row_cells):
                row_cells[col_idx].text = val
                set_cell_background(row_cells[col_idx], bg_color)
                
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                
                if p.runs:
                    run = p.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)
                    # Make OMP and BVH bold if helpful, or just standard
                    if col_idx == 0:
                        run.bold = True

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)

if __name__ == "__main__":
    parse_markdown_to_docx("rapor.md", "rapor.docx")
